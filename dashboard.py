import streamlit as st
import pandas as pd
import json
import os
import zipfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from fpdf import FPDF
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import openai
from datetime import datetime
import re

# Set your OpenAI API key
openai.api_key = 'sk-proj-RzYho-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'

# Function to generate a report using the ChatGPT API
def generate_report_with_chatgpt(data):
    try:
        prompt = (
            "You are a data analyst that has to visualize data as provided. You have to write a formal report that will be submitted to the company's client. The record contains 7 structure . One is cover page as 1 page, second is abstract as 1 page and title of abstract should be in the middle. "
            "Third one is the introduction that provides overview of data. The overview should counted every participants of survey, province, age, what crops they mostly plant and other variables. The fourth one is Data visualization based on data what can we see. "
            "So You have to make graphs to explain . Moreover, the outcome focus more on the average of total yield of each crops compared to land area that they use for plants. For example, in 2024 they planted corn in 200m square and they got 1000kg and in 6 months after project they plant 200m but I got 20000kg. So you have to report this changes"
            "Be careful that some vegetables weren't planted in 2024 so don't compare it just described it briefly. The fifth part is results based on data visualization and output. Last but not least, you must conclusion every thing in the conclusion part along with recommandations."
            f"{json.dumps(data, indent=2)}"
        )
        
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",  # Or "gpt-3.5-turbo"
            messages=[{"role": "system", "content": "You are a helpful assistant."}, {"role": "user", "content": prompt}],
            max_tokens=8000,
            temperature=0.7
        )
        
        report_parts = response.choices[0].message['content'].strip()
        st.write(report_parts)
        
        # Save report as Word and PDF documents
        word_filename = 'report.docx'
        pdf_filename = 'report.pdf'
        save_report_as_word(report_parts, word_filename)
        save_report_as_pdf(report_parts, pdf_filename)
        
        return report_parts, word_filename, pdf_filename
    except Exception as e:
        st.error(f"Failed to generate report: {e}")
        return None, None, None


def create_cover_page(doc):
    # Create cover page
    section = doc.sections[0]

    # Add blank paragraphs to push the title to the center of the page
    for _ in range(10):
        doc.add_paragraph("\n")
        
    # Add title
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(f'Six Month Progress Report \n{datetime.now().strftime("%B %Y")}')
    run.font.size = Pt(20)
    run.bold = True

    # Add blank paragraphs after the title to fill the rest of the page
    for _ in range(10):
        doc.add_paragraph("\n")

    # Ensure the cover page is a separate section
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)

def add_formatted_text(paragraph, text):
    """
    Add text to a paragraph with formatting based on markdown-like syntax.
    """
    # Pattern to match **bold**
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    cursor = 0
    for match in bold_pattern.finditer(text):
        # Add text before the bold pattern
        if cursor < match.start():
            paragraph.add_run(text[cursor:match.start()])

        # Add bold text
        paragraph.add_run(match.group(1)).bold = True
        cursor = match.end()

    # Add remaining text after the last match
    if cursor < len(text):
        paragraph.add_run(text[cursor:])

def save_report_as_word(report, filename):
    try:
        doc = Document()
        
        # Create cover page
        create_cover_page(doc)

        # Content
        lines = report.split('\n')
        table = None
        for line in lines:
            if line.strip().startswith("# "):
                paragraph = doc.add_heading(line.strip()[2:], level=1)
            elif line.strip().startswith("## "):
                paragraph = doc.add_heading(line.strip()[3:], level=2)
            elif line.strip().startswith("### "):
                paragraph = doc.add_heading(line.strip()[4:], level=3)
            elif line.strip().startswith("* "):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif "|" in line:
                table_data = [cell.strip() for cell in line.split('|') if cell]
                if not table:
                    table = doc.add_table(rows=1, cols=len(table_data))
                    hdr_cells = table.rows[0].cells
                    for i, cell_data in enumerate(table_data):
                        hdr_cells[i].text = cell_data
                else:
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(table_data):
                        row_cells[i].text = cell_data
            else:
                paragraph = doc.add_paragraph()
                add_formatted_text(paragraph, line)

        doc.save(filename)
    except Exception as e:
        st.error(f"Failed to save Word report: {e}")


# Enhanced PDF class for better formatting
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Report', 0, 1, 'C')
        self.set_font('Arial', '', 10)
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title, 0, 1, 'L')
        self.ln(10)

    def chapter_body(self, body):
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, body)
        self.ln()

def save_report_as_pdf(report, filename):
    try:
        pdf = PDF()
        pdf.add_page()

        # Cover Page
        pdf.set_font('Arial', 'B', 20)
        pdf.cell(0, 10, 'Six Month Progress Report', 0, 1, 'C')
        pdf.cell(0, 10, datetime.now().strftime('%B %Y'), 0, 1, 'C')
        pdf.add_page()

        # Content
        lines = report.split('\n')
        for line in lines:
            if line.strip().startswith("# "):
                pdf.chapter_title(line.strip()[2:])
            elif line.strip().startswith("## "):
                pdf.set_font('Arial', 'B', 12)
                pdf.multi_cell(0, 10, line.strip()[3:])
                pdf.ln()
            elif line.strip().startswith("### "):
                pdf.set_font('Arial', 'B', 12)
                pdf.multi_cell(0, 10, line.strip()[4:])
                pdf.ln()
            elif line.strip().startswith("* "):
                pdf.set_font('Arial', '', 12)
                pdf.cell(0, 10, '\u2022 ' + line.strip()[2:], 0, 1)
            elif "|" in line:
                table_data = [cell.strip() for cell in line.split('|') if cell]
                col_width = pdf.w / len(table_data) - 1
                for cell_data in table_data:
                    pdf.cell(col_width, 10, cell_data, 1)
                pdf.ln(10)
            else:
                pdf.chapter_body(line)

        pdf.output(filename)
    except Exception as e:
        st.error(f"Failed to save PDF report: {e}")

# Function to create a zip file containing the Word and PDF reports
def create_zip_file(word_filename, pdf_filename, zip_filename):
    try:
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            zipf.write(word_filename)
            zipf.write(pdf_filename)
        st.success(f"Zip file {zip_filename} created successfully.")
    except Exception as e:
        st.error(f"Failed to create zip file: {e}")

# Function to send email with attachments
def send_email_with_attachments(subject, body, attachments):
    from_email = "seaklav168@gmail.com"
    password = "wnfj ptne wqhf joie"
    to_email = ["hratana261@gmail.com" ,"khengdalish21@gmail.com","chlakhna702@gmail.com"] # Replace with the fixed email address

    # Create the email message
    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = ", ".join(to_email)
    msg['Subject'] = subject

    # Attach the body text
    msg.attach(MIMEText(body, 'plain'))

    # Attach the files
    for attachment in attachments:
        try:
            with open(attachment, "rb") as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename={attachment}')
                msg.attach(part)
        except Exception as e:
            st.error(f"Failed to attach file {attachment}: {e}")

    # Send the email
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
        st.success(f"Email sent to {to_email}")
    except Exception as e:
        st.error(f"Failed to send email: {e}")

# Function to fetch data from Google Sheets
def fetch_data(google_sheet_url):
    try:
        df = pd.read_csv(google_sheet_url)
    except Exception as e:
        st.error(f"Failed to fetch data from Google Sheets: {e}")
        return None
    return df

# Function to fetch pivot table data
def fetch_pivot_data(pivot_table_url):
    try:
        pivot_df = pd.read_csv(pivot_table_url)
    except Exception as e:
        st.error(f"Failed to fetch pivot table data from Google Sheets: {e}")
        return None
    return pivot_df

# Function to render the dashboard
def dashboard():
    st.set_page_config(
        page_title="DCx Co.,ltd",
        page_icon="https://dcxsea.com/asset/images/logo/LOGO_DCX.png",
        layout="wide",
        initial_sidebar_state="collapsed"
    )

    hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                </style>
                """
    st.markdown(hide_st_style, unsafe_allow_html=True)

    st.markdown(
        """
        <div style="display: flex; align-items: center;">
            <img src="https://cdn3d.iconscout.com/3d/free/thumb/free-line-chart-growth-3814121-3187502.png" alt="logo" style="width: 90px; margin-right: 15px;">
            <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0;">ការបន្សាំកសិកម្មជនជាតិដើមភាគតិច</h3>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.markdown(
        """
        <div style="display: flex; justify-content: center;margin-top: 0px; margin-bottom: 20px;">
            <img src="https://dcxsea.com/asset/images/logo/LOGO_DCX.png" style="width: 150px;">
        </div>
        """,
        unsafe_allow_html=True
    )

    options = st.sidebar.selectbox(
        'Choose Dataset',
        [' ', '6 Months', 'One Year', '6 & 12 Months']
    )

    if options == 'One Year':
        df = fetch_data(
            google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=1474899771&single=true&output=csv'
        )
    elif options == '6 Months':
        df = fetch_data(
            google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=0&single=true&output=csv'
        )
    elif options == '6 & 12 Months':
        df = fetch_data(
            google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=0&single=true&output=csv'
        )
    else:
        # Fetch and display the pivot table
        st.markdown(
            """
            <div style="display: flex; align-items: center;">
                <img src="https://symbolshub.org/wp-content/uploads/2019/10/bullet-point-symbol.png" alt="logo" style="width: 25px; margin-right: 5px; vertical-align: middle;">
                <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0; font-size: 18px; font-weight: bold; vertical-align: middle;">ទិន្នន័យអំពីការបន្សាំកសិកម្មនៃជនជាតិភាគតិច</h3><br><br><br>
            </div>
            """,
            unsafe_allow_html=True
        )
        pivot_table_url = 'https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=254021688&single=true&output=csv'
        pivot_df = fetch_pivot_data(pivot_table_url)
        if pivot_df is not None:
            pivot = pivot_df.style.set_properties(**{'background-color': 'rgb(161, 219, 255, 0.3)', 'color': 'white'})
            st.dataframe(pivot)

    if options in ['One Year', '6 Months', '6 & 12 Months'] and df is not None:
        st.markdown(
            """
            <div style="display: flex; align-items: center;">
                <img src="https://symbolshub.org/wp-content/uploads/2019/10/bullet-point-symbol.png" alt="logo" style="width: 25px; margin-right: 5px; vertical-align: middle;">
                <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0; font-size: 18px; font-weight: bold; vertical-align: middle;">ទិន្នន័យអំពីការបន្សាំកសិកម្មនៃជនជាតិភាគតិច</h3><br><br><br>
            </div>
            """,
            unsafe_allow_html=True
        )
        df_style = df.style.set_properties(**{'background-color': 'rgb(161, 219, 255, 0.3)', 'color': 'white'})
        st.dataframe(df_style)

        # Add button to generate report with custom style
        button_style = """
            <style>
            .stButton>button {
                background-color: green;
                color: white;
                border: none;
                padding: 0.5em 1em;
                cursor: pointer;
            }
            .stButton>button:hover {
                background-color: darkgreen;
                color: #DFFF00;
            }
            </style>
        """
        st.markdown(button_style, unsafe_allow_html=True)

        if st.button('Generate Report'):
            if df.empty:
                st.error("No data available to send.")
            else:
                df_cleaned = df.fillna('')  # Fill NaNs with empty strings
                data = df_cleaned.to_dict(orient='records')  # Convert DataFrame to dictionary format

                # Generate report with the ChatGPT API
                report_content, word_filename, pdf_filename = generate_report_with_chatgpt(data)

                if report_content:
                    zip_filename = 'reports.zip'
                    create_zip_file(word_filename, pdf_filename, zip_filename)
                    
                    # Send email with the zip file
                    send_email_with_attachments("Generated Report", "Please find the attached reports.", [zip_filename])
                    
                    # Allow downloading of the zip file
                    st.download_button('Download All Reports', data=open(zip_filename, 'rb').read(), file_name=zip_filename, mime='application/zip')
                else:
                    st.write("Failed to generate report.")

if __name__ == '__main__':
    dashboard()
